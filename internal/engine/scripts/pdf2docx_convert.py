import sys
import re
from pdf2docx import Converter
import fitz
from docx import Document

src, dst = sys.argv[1], sys.argv[2]


def convert_with_pdf2docx(src_path, dst_path):
    cv = Converter(src_path)
    cv.convert(dst_path)
    cv.close()


def docx_text(path):
    d = Document(path)
    chunks = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    return '\n'.join(chunks)


def output_looks_broken(text):
    if not text:
        return False
    letters = sum(1 for ch in text if ch.isalpha())
    spaces = text.count(' ')
    ratio = spaces / max(letters, 1)
    run_on = sum(1 for tok in re.split(r'\s+', text) if len(tok) >= 35)
    return ratio < 0.045 or run_on >= 10


def words_to_lines(words):
    words.sort(key=lambda w: ((w[1] + w[3]) / 2.0, w[0]))
    lines = []
    current = []
    current_y = None
    for w in words:
        x0, y0, x1, y1, token = w[:5]
        y = (y0 + y1) / 2.0
        if current_y is None or abs(y - current_y) <= 2.5:
            current.append((x0, x1, token))
            if current_y is None:
                current_y = y
            else:
                current_y = (current_y + y) / 2.0
        else:
            if current:
                current.sort(key=lambda t: t[0])
                lines.append(current)
            current = [(x0, x1, token)]
            current_y = y
    if current:
        current.sort(key=lambda t: t[0])
        lines.append(current)
    return lines


def split_cells(line, gap_threshold=95.0):
    if not line:
        return []
    cells = []
    current = [line[0][2]]
    prev_x1 = line[0][1]
    for x0, x1, token in line[1:]:
        if x0 - prev_x1 >= gap_threshold:
            cells.append(' '.join(current).strip())
            current = [token]
        else:
            current.append(token)
        prev_x1 = x1
    cells.append(' '.join(current).strip())
    return [c for c in cells if c]


def line_text(line):
    return ' '.join(token for _, _, token in line).strip()


def is_heading(text):
    if not text:
        return False
    if re.match(r'^\d+(\.\d+)*\s+\S+', text):
        return True
    words = text.split()
    if len(words) <= 8 and text == text.title() and not text.endswith(('.', ';', ',')):
        return True
    return False


def emit_paragraph(doc, text):
    if not text:
        return
    if is_heading(text):
        doc.add_heading(text, level=2)
    else:
        doc.add_paragraph(text)


def emit_table(doc, rows):
    if len(rows) < 2:
        return False
    width = max(len(r) for r in rows)
    if width < 3:
        return False
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(width):
            table.cell(i, j).text = row[j] if j < len(row) else ''
    return True


def rebuild_readable_docx(src_path, dst_path):
    pdf = fitz.open(src_path)
    doc = Document()
    added = 0
    for page in pdf:
        words = page.get_text('words')
        if words:
            lines = words_to_lines(words)
            i = 0
            while i < len(lines):
                row = split_cells(lines[i])
                if len(row) >= 3:
                    rows = [row]
                    j = i + 1
                    while j < len(lines):
                        nxt = split_cells(lines[j])
                        if len(nxt) >= 3:
                            rows.append(nxt)
                            j += 1
                        else:
                            break
                    if emit_table(doc, rows):
                        added += len(rows)
                        i = j
                        continue
                text = line_text(lines[i])
                if text:
                    emit_paragraph(doc, text)
                    added += 1
                i += 1
        else:
            plain = page.get_text('text')
            for raw in plain.splitlines():
                text = raw.strip()
                if text:
                    emit_paragraph(doc, text)
                    added += 1
        if page.number < len(pdf) - 1:
            doc.add_paragraph('')

    pdf.close()
    if added > 0:
        doc.save(dst_path)


convert_with_pdf2docx(src, dst)
text = docx_text(dst)
if output_looks_broken(text):
    rebuild_readable_docx(src, dst)
